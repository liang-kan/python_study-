import os
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches


# --- 模块1：模拟 BIM 引擎 (The "Fake" BIM Engine) ---
# 在真实项目中，这里会替换为调用 Autodesk Forge 或 BIMFACE 的截图 API
def mock_bim_renderer(component_type, filename):
    """
    根据构件类型，动态绘制一张示意图并保存。
    """
    plt.figure(figsize=(4, 3))
    plt.title(f"BIM Model View: {component_type}", fontsize=14)
    plt.axis('off')  # 关闭坐标轴

    # 模拟不同构件的“渲染”
    if component_type == "混凝土柱":
        # 画一个圆代表柱子
        circle = plt.Circle((0.5, 0.5), 0.3, color='grey', alpha=0.8)
        plt.gca().add_patch(circle)
        plt.text(0.5, 0.5, "Column (Z-1)", ha='center', color='white')

    elif component_type == "钢筋混凝土梁":
        # 画一个矩形代表梁
        rect = plt.Rectangle((0.1, 0.4), 0.8, 0.2, color='blue', alpha=0.6)
        plt.gca().add_patch(rect)
        plt.text(0.5, 0.5, "Beam (L-2)", ha='center', color='white')

    elif component_type == "底板基础":
        # 画一大块区域代表基础
        rect = plt.Rectangle((0, 0), 1, 0.8, color='brown', alpha=0.5)
        plt.gca().add_patch(rect)
        plt.text(0.5, 0.4, "Foundation", ha='center', color='white')

    # 保存图片
    plt.savefig(filename, bbox_inches='tight')
    plt.close()
    return filename


# --- 模块2：逻辑核心 (The Logic Core) ---

def generate_bid_document(article_paragraphs, output_filename):
    doc = Document()
    doc.add_heading('自动化工程标书示例', 0)

    # 定义关键词与BIM构件的映射关系
    # 真实场景中，这里会使用 NLP 模型进行语义匹配
    keywords_map = {
        "柱": "混凝土柱",
        "梁": "钢筋混凝土梁",
        "基础": "底板基础",
        "底板": "底板基础"
    }

    print("开始处理文章...")

    for text in article_paragraphs:
        # 1. 写入文字段落
        p = doc.add_paragraph(text)

        # 2. 分析这段文字，看是否包含关键词
        matched_component = None
        for key, component_name in keywords_map.items():
            if key in text:
                matched_component = component_name
                break  # 简单起见，匹配到一个就停止

        # 3. 如果匹配到了，模拟“去BIM模型里截图”并插入
        if matched_component:
            print(f" -> 发现关键词，正在生成 [{matched_component}] 的BIM图...")

            # 临时图片路径
            img_path = f"temp_{matched_component}.png"

            # 调用“BIM引擎”生成图片
            mock_bim_renderer(matched_component, img_path)

            # 插入图片到 Word
            doc.add_picture(img_path, width=Inches(4.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # 图片居中

            # 既然是Demo，用完就把临时图片删了（可选）
            # os.remove(img_path)

    doc.save(output_filename)
    print(f"\n成功！标书已生成: {output_filename}")


# --- 模块3：运行入口 ---

if __name__ == "__main__":
    # 模拟 AI 生成的标书段落内容
    dummy_article = [
        "第一章：工程概况",
        "本工程位于市中心，总建筑面积5万平方米。",
        "第二章：主要施工工艺",
        "2.1 基础工程施工",
        "在进行地下室施工时，首先要进行大体积混凝土浇筑。底板基础的厚度达到了1500mm，需要严格控制水化热。",
        "2.2 竖向构件施工",
        "首层大厅采用C60高强混凝土柱，截面尺寸为800x800mm，施工时需确保垂直度。",
        "2.3 水平构件施工",
        "二层结构中，主跨位置设置了预应力钢筋混凝土梁，跨度较大，需提前安装波纹管。"
    ]

    generate_bid_document(dummy_article, "Project_Bid_Demo.docx")