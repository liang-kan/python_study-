import streamlit as st
import requests
import json
import time
from docx import Document
from io import BytesIO


# --- 1. 核心逻辑函数 (复用之前的) ---

def ask_ai(prompt, api_key):
    """AI 请求函数"""
    url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    data = {
        "model": "glm-4",
        "messages": [
            {"role": "system", "content": "你是一个投标专家。请用 Markdown 格式输出。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"Error: {response.text}"
    except Exception as e:
        return str(e)


def write_markdown_to_docx(doc, text):
    """Markdown 转 Word 简单版"""
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)


# --- 2. Streamlit 界面代码 ---

# 页面标题配置
st.set_page_config(page_title="AI 标书生成神器", page_icon="📝")

st.title("🚀 AI 标书生成神器")
st.write("上传招标需求，输入公司优势，一键生成专业投标文件！")

# 侧边栏：配置区
with st.sidebar:
    st.header("⚙️ 设置")
    api_key = st.text_input("请输入智谱 API Key", type="password")

    st.info("提示：这只是一个原型演示。请确保 API Key 有效。")

# 主区域：输入区
col1, col2 = st.columns(2)

with col1:
    st.subheader("1. 招标需求")
    # 这里简化为直接输入文本，真实场景可以是上传 PDF 并解析
    bid_requirement = st.text_area("粘贴招标文件的核心需求：", height=200,
                                   placeholder="例如：本项目采购100台电脑，预算50万...")

with col2:
    st.subheader("2. 公司优势 (知识库)")
    company_info = st.text_area("粘贴公司介绍/案例/资质：", height=200,
                                placeholder="例如：我司成立于2015年，拥有ISO认证...")

# 章节设置
st.subheader("3. 目录规划")
default_chapters = "第一章 公司简介\n第二章 技术方案\n第三章 实施计划\n第四章 售后服务"
chapters_input = st.text_area("每行一个章节标题：", value=default_chapters, height=150)

# 生成按钮
if st.button("🔥 开始生成标书", type="primary"):
    if not api_key:
        st.error("请先在左侧输入 API Key！")
    elif not bid_requirement:
        st.warning("请输入招标需求！")
    else:
        # --- 开始生成流程 ---

        # 显示进度条
        progress_bar = st.progress(0)
        status_text = st.empty()

        # 准备 Word 对象
        doc = Document()
        doc.add_heading('投标文件', 0)

        # 解析章节
        chapters = [line.strip() for line in chapters_input.split('\n') if line.strip()]
        total_chapters = len(chapters)

        # 存放生成结果的容器 (用于在网页展示预览)
        preview_container = st.container()

        for i, chapter in enumerate(chapters):
            # 更新状态
            status_text.text(f"正在撰写：{chapter} ...")

            doc.add_heading(chapter, level=1)

            # 构建 Prompt
            prompt = f"""
            【招标需求】{bid_requirement}
            【公司资料】{company_info}
            【任务】撰写【{chapter}】章节。
            【要求】Markdown格式，分点论述，引用公司资料，字数300字。
            """

            # 调用 AI
            content = ask_ai(prompt, api_key)

            # 写入 Word
            write_markdown_to_docx(doc, content)

            # 在网页上实时预览
            with preview_container:
                with st.expander(f"✅ {chapter} (点击展开预览)", expanded=False):
                    st.markdown(content)

            # 更新进度条
            progress_bar.progress((i + 1) / total_chapters)

        status_text.success("🎉 生成完成！")

        # --- 提供下载 ---
        # 将 doc 对象保存到内存流中，而不是硬盘文件
        # 这样用户点击下载时，是从内存直接下载，适合 Web 环境
        bio = BytesIO()
        doc.save(bio)

        st.download_button(
            label="📥 下载 Word 文档 (.docx)",
            data=bio.getvalue(),
            file_name="AI生成的投标文件.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )