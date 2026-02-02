import os
import time
from zhipuai import ZhipuAI
from docx import Document
from docx.shared import Pt

# --- 配置 ---
API_KEY = "8354280bc4d84b13aea514cd0b55c8a2.97vqKhN5Nlw1MvBV"
client = ZhipuAI(api_key=API_KEY)


# 读取之前的需求分析结果
def load_analysis():
    if os.path.exists("bid_analysis.txt"):
        with open("bid_analysis.txt", "r", encoding="utf-8") as f:
            return f.read()
    return "（未找到需求分析文件，请先运行上一课的代码）"


# 通用 AI 调用函数 (非流式，一次性返回)
def ask_ai(prompt):
    try:
        response = client.chat.completions.create(
            model="glm-4",
            messages=[
                {"role": "system", "content": "你是一个专业的投标专员，擅长撰写严谨、详实的投标文件。"},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"❌ AI 调用出错: {e}")
        return ""


# --- 核心流程 ---
def main():
    print("🚀 标书生成引擎启动...")

    # 1. 获取背景信息
    analysis_text = load_analysis()
    print(f"📖 已加载需求分析，长度: {len(analysis_text)} 字")

    # 2. 生成目录大纲
    print("\n📋 正在规划目录大纲...")
    outline_prompt = f"""
    基于以下招标需求分析，请为我设计一份投标文件的【一级目录】大纲。
    需求分析：
    {analysis_text[:2000]}

    要求：
    1. 只返回目录标题，每行一个。
    2. 不要有多余的废话。
    3. 必须包含：项目理解、技术方案、实施计划、售后服务、资质证明。
    4. 格式示例：
    第一章 项目背景与理解
    第二章 核心技术方案
    ...
    """

    outline_raw = ask_ai(outline_prompt)

    # 清洗目录 (把空行去掉)
    chapters = [line.strip() for line in outline_raw.split('\n') if line.strip()]

    print("-" * 30)
    print("✅ AI 规划的目录:")
    for ch in chapters:
        print(ch)
    print("-" * 30)

    confirm = input("👉 目录是否满意？(y/n): ")
    if confirm.lower() != 'y':
        print("请修改代码中的 prompt 重新生成，或手动修改 chapters 列表。")
        return

    # 3. 创建 Word 文档
    doc = Document()
    doc.add_heading('投标文件', 0)

    # 4. 逐章生成内容 (这是最精彩的部分)
    print("\n✍️ 开始逐章撰写正文...")

    for chapter_title in chapters:
        print(f"   ➤ 正在写: {chapter_title} ...", end="", flush=True)

        # 将标题写入 Word
        doc.add_heading(chapter_title, level=1)

        # 构建这一章的 Prompt
        content_prompt = f"""
        背景：{analysis_text[:1000]}

        当前任务：请为投标文件的【{chapter_title}】章节撰写详细正文。
        要求：
        1. 内容要专业、详实，符合招标书要求。
        2. 既然是{chapter_title}，就要侧重于该方面的描述。
        3. 字数控制在 300-500 字左右。
        4. 直接输出正文，不要包含“好的”、“如下”等废话。
        """

        # 调用 AI
        chapter_content = ask_ai(content_prompt)

        # 写入 Word
        doc.add_paragraph(chapter_content)

        print(" ✅ 完成")

        # 休息一下，防止 API 速率限制
        time.sleep(2)

    # 5. 保存最终文件
    output_filename = "最终投标文件.docx"
    doc.save(output_filename)
    print(f"\n🎉 大功告成！文件已保存为: {output_filename}")


if __name__ == "__main__":
    main()