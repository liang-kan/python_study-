import os
import time
import requests
import json
from docx import Document

# --- 配置 ---
API_KEY = "8354280bc4d84b13aea514cd0b55c8a2.97vqKhN5Nlw1MvBV"


# 1. 读取公司知识库
def load_company_profile():
    if os.path.exists("company_profile.txt"):
        with open("company_profile.txt", "r", encoding="utf-8") as f:
            return f.read()
    return "（无公司资料，请自行补充）"


# 读取需求分析
def load_analysis():
    if os.path.exists("bid_analysis.txt"):
        with open("bid_analysis.txt", "r", encoding="utf-8") as f:
            return f.read()
    return "（需求分析为空）"


# AI 请求函数
def ask_ai(prompt):
    url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {API_KEY}"
    }
    data = {
        "model": "glm-4",
        "messages": [
            {"role": "system",
             "content": "你是一个资深的投标专家。你的写作风格应当专业、自信，并善于将公司的优势与招标需求结合起来。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        result = response.json()
        return result['choices'][0]['message']['content']
    except Exception as e:
        print(f"❌ 请求异常: {e}")
        return ""


def main():
    print("🚀 标书生成引擎 PRO 版启动...")

    analysis_text = load_analysis()
    company_info = load_company_profile()  # 加载公司资料

    print(f"📖 已加载需求分析: {len(analysis_text)} 字")
    print(f"🏢 已加载公司资料: {len(company_info)} 字")

    # --- 阶段1：生成目录 ---
    print("\n📋 正在规划目录大纲...")
    outline_prompt = f"""
    基于以下招标需求：
    {analysis_text[:1000]}

    请设计一份投标文件的【一级目录】。
    要求：包含项目理解、技术方案、实施计划、售后服务、公司优势与案例。
    只返回目录标题，每行一个。
    """
    outline_raw = ask_ai(outline_prompt)
    chapters = [line.strip() for line in outline_raw.split('\n') if line.strip()]

    print("✅ 目录规划完成。")

    # --- 阶段2：创建文档 ---
    doc = Document()
    doc.add_heading('投标文件', 0)

    print("\n✍️ 开始逐章撰写 (已开启知识库注入)...")

    for chapter_title in chapters:
        print(f"   ➤ 正在写: {chapter_title} ...", end="", flush=True)
        doc.add_heading(chapter_title, level=1)

        # --- 关键修改：Prompt Engineering ---
        # 我们告诉 AI：
        # 1. 这一章的标题是什么
        # 2. 招标方的需求是什么
        # 3. 我们公司有什么优势 (Knowledge Injection)
        # 4. 必须把这两者结合起来写 (Instruction)

        content_prompt = f"""
        【背景信息】
        招标需求摘要：{analysis_text[:800]}

        【我方公司资料】
        {company_info}

        【当前任务】
        请撰写投标文件的【{chapter_title}】章节。

        【写作要求】
        1. **必须引用我方公司资料中的具体数据或案例**来证明我们能满足招标需求。
           (例如：如果写技术方案，要提到我们的某某产品；如果写售后，要提到我们的响应时间)
        2. 如果公司资料里没有相关内容，可以适当进行通用的专业描述，但不要瞎编数据。
        3. 语气要诚恳且自信。
        4. 字数约 400 字。
        5. 不要写“根据公司资料显示”这种话，直接以第一人称“我公司”来写。
        """

        chapter_content = ask_ai(content_prompt)
        doc.add_paragraph(chapter_content)

        print(" ✅ 完成")
        time.sleep(1)

    doc.save("精准投标文件.docx")
    print(f"\n🎉 大功告成！文件已保存。")


if __name__ == "__main__":
    main()