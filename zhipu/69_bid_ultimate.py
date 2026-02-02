import os
import time
import requests
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# --- 配置 ---
API_KEY = "8354280bc4d84b13aea514cd0b55c8a2.97vqKhN5Nlw1MvBV"

# --- 1. 结构化公司知识库 (模拟数据库) ---
# 键：关键词列表；值：对应的资料内容
COMPANY_DB = {
    ("技术", "方案", "产品", "功能"): """
    【我方技术优势】
    1. 核心产品：'天眼'智能监控系统，支持 4K 实时流分析。
    2. 专利算法：自研 'EagleEye' 图像识别引擎，准确率 99.8%。
    3. 架构：采用微服务架构，支持 Docker 容器化部署，弹性伸缩。
    """,

    ("售后", "服务", "培训", "维护"): """
    【我方服务承诺】
    1. 质保期：提供 3 年免费质保（优于国标）。
    2. 响应速度：7*24小时客服，15分钟内响应，2小时内到达现场。
    3. 培训计划：提供每季度一次的上门巡检和技术培训。
    """,

    ("案例", "业绩", "经验", "证明"): """
    【我方成功案例】
    1. 某省公安厅项目：金额 5000 万，已稳定运行 3 年。
    2. 某大型机场安防项目：覆盖 10 万平米，零误报。
    """,

    ("资质", "简介", "概况", "背景"): """
    【公司概况】
    未来科技有限公司，成立于2015年，注册资金1亿。
    资质：CMMI5, ISO9001, 国家高新技术企业, 涉密资质乙级。
    """
}


# --- 2. 智能检索函数 (RAG 雏形) ---
def get_relevant_info(chapter_title):
    """
    根据章节标题，去知识库里找最相关的内容
    """
    context = ""
    print(f"   🔍 正在为【{chapter_title}】检索资料...", end="")

    matched = False
    for keywords, content in COMPANY_DB.items():
        # 只要标题里包含任意一个关键词
        for kw in keywords:
            if kw in chapter_title:
                context += content + "\n"
                matched = True
                break  # 命中一个关键词就够了，不用重复添加

    if matched:
        print(" ✅ 命中知识库")
        return context
    else:
        print(" ⚠️ 未命中 (使用通用资料)")
        # 如果没匹配到，返回公司简介作为保底
        return COMPANY_DB[("资质", "简介", "概况", "背景")]


# --- AI 请求 ---
def ask_ai(prompt):
    url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}
    data = {
        "model": "glm-4",
        "messages": [
            {"role": "system", "content": "你是一个投标专家。请用 Markdown 格式输出。"},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        return response.json()['choices'][0]['message']['content']
    except:
        return ""


# --- 3. Markdown 转 Word 排版引擎 ---
def write_markdown_to_docx(doc, text):
    """
    简单的解析器：把 Markdown 文本转换成 Word 的格式
    支持：
    # 二级标题
    ## 三级标题
    - 列表项
    **加粗** (简单处理)
    """
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            # 列表项
            p = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            # 普通段落
            p = doc.add_paragraph(line)
            # 简单的字体设置
            p.paragraph_format.space_after = Pt(6)


def main():
    print("🚀 标书生成引擎 Ultimate 版启动...")

    # 假设这是之前分析出来的需求（为了演示方便直接写死）
    analysis_text = "本项目采购一套智能安防系统，要求支持人脸识别，预算200万，工期3个月。"

    # 假设这是目录
    chapters = ["第一章 公司简介与资质", "第二章 核心技术方案", "第三章 成功案例展示", "第四章 售后服务承诺"]

    doc = Document()
    # 设置大标题
    title = doc.add_heading('投标文件', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    print("\n✍️ 开始生成...")

    for chapter_title in chapters:
        print(f"➤ 章节: {chapter_title}")

        # 1. 智能找资料
        relevant_info = get_relevant_info(chapter_title)

        # 2. 写入一级标题
        doc.add_heading(chapter_title, level=1)

        # 3. 构造 Prompt
        content_prompt = f"""
        【背景】招标需求：{analysis_text}
        【参考资料】{relevant_info}

        【任务】请撰写【{chapter_title}】的内容。

        【重要格式要求】
        1. 请使用 Markdown 格式。
        2. 使用 ## 表示小标题（不要用 #，因为 # 是一级标题）。
        3. 使用 - 表示列表。
        4. 必须引用【参考资料】中的数据。
        5. 字数 300 字左右。
        """

        # 4. AI 生成
        md_content = ask_ai(content_prompt)

        # 5. 解析并写入 Word
        write_markdown_to_docx(doc, md_content)

        print("   ✅ 写入完成\n")
        time.sleep(1)

    doc.save("完美排版投标文件.docx")
    print(f"🎉 文件已保存！")


if __name__ == "__main__":
    main()