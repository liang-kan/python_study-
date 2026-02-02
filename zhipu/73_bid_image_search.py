import requests
import time
from docx import Document
from docx.shared import Inches
from io import BytesIO
from duckduckgo_search import DDGS  # 搜图神器


# 1. 搜图函数
def search_and_download_image(keyword, max_retries=3):
    """
    根据关键词搜索图片，并返回二进制图片流
    """
    print(f"   🔍 正在全网搜索图片：【{keyword}】...", end="")

    try:
        # 使用 DuckDuckGo 搜索图片
        with DDGS() as ddgs:
            # 搜索关键词，safesearch='off' 表示不过滤，max_results=3 只要前3张
            results = list(ddgs.images(keyword, max_results=5))

        if not results:
            print(" ❌ 未找到相关图片")
            return None

        # 尝试下载图片（因为有些链接可能防盗链或打不开，所以要试几个）
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        }

        for img_data in results:
            img_url = img_data['image']
            try:
                # 设置超时，防止卡死
                resp = requests.get(img_url, headers=headers, timeout=5)
                if resp.status_code == 200:
                    print(" ✅ 下载成功")
                    return BytesIO(resp.content)
            except:
                continue  # 这张下不了，试下一张

        print(" ❌ 所有图片下载失败")
        return None

    except Exception as e:
        print(f" ❌ 搜索出错: {e}")
        return None


# --- 主流程 ---
def main():
    print("🌍 标书配图引擎启动...")

    doc = Document()
    doc.add_heading('设备选型与参数', 0)

    # 假设这是从 AI 生成的标书里提取出来的设备清单
    # (真实场景下，你可以让 AI 输出一个 JSON 列表)
    equipment_list = [
        {"name": "戴尔 PowerEdge R750 服务器", "desc": "高性能计算节点，支持双路 CPU。"},
        {"name": "海康威视 400万像素 摄像头", "desc": "红外夜视，支持 POE 供电。"},
        {"name": "华为 S5720 交换机", "desc": "千兆接入，万兆上行。"}
    ]

    for item in equipment_list:
        name = item['name']
        desc = item['desc']

        # 1. 写入文字介绍
        doc.add_heading(name, level=2)
        doc.add_paragraph(f"产品描述：{desc}")

        # 2. 搜索并插入图片
        # 技巧：关键词加上 "产品图" 或 "白底"，搜出来的图更像标书用的
        search_keyword = name + " 产品图"

        img_stream = search_and_download_image(search_keyword)

        if img_stream:
            try:
                # 插入图片，宽度固定为 4 英寸，保持整齐
                doc.add_picture(img_stream, width=Inches(4))
                doc.add_paragraph(f"图：{name} 实物参考图", style="Caption")
            except Exception as e:
                print(f"   ⚠️ 图片格式 Word 不支持，跳过。({e})")

        print("-" * 20)
        time.sleep(1)  # 礼貌爬取

    doc.save("配图版标书.docx")
    print("\n🎉 文件已保存！")


if __name__ == "__main__":
    main()