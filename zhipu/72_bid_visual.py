import requests
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO

# --- 配置 ---
API_KEY = "你的_API_KEY_在这里"  # 智谱 API


# 1. AI 生成 Mermaid 代码
def generate_diagram_code(requirement, diagram_type="flowchart"):
    """
    让 AI 根据需求生成 mermaid 代码
    diagram_type: flowchart (流程图) / gantt (甘特图)
    """
    url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {API_KEY}"}

    prompt = f"""
    请根据以下项目需求，生成一个 {diagram_type} 类型的 Mermaid 图表代码。
    需求：{requirement}

    【重要要求】
    1. 只返回 Mermaid 代码，不要包含 ```mermaid 或其他废话。
    2. 如果是流程图(flowchart)，请用 graph TD 开头，节点要有中文描述。
    3. 如果是甘特图(gantt)，请规划合理的阶段（需求、开发、测试、部署），注意日期格式。
    4. 代码必须语法正确，能够直接被渲染。
    """

    data = {
        "model": "glm-4",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.1  # 低温，保证代码准确
    }

    try:
        res = requests.post(url, headers=headers, json=data)
        code = res.json()['choices'][0]['message']['content']
        # 清洗代码 (去掉可能存在的 markdown 标记)
        code = code.replace("```mermaid", "").replace("```", "").strip()
        return code
    except:
        return None


# 2. 将 Mermaid 代码转为图片 (使用 mermaid.ink 在线服务)
def mermaid_to_image(mermaid_code):
    """
    利用 mermaid.ink 的公共 API 把代码转成图片流
    """
    # 需要对代码进行 base64 编码
    graphbytes = mermaid_code.encode("utf8")
    base64_bytes = base64.urlsafe_b64encode(graphbytes)
    base64_string = base64_bytes.decode("ascii")

    # 构造请求 URL
    img_url = "https://mermaid.ink/img/" + base64_string

    print(f"   🖼️ 图表生成地址: {img_url}")

    try:
        response = requests.get(img_url)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print("❌ 图片下载失败，可能是 Mermaid 代码语法错误。")
            return None
    except Exception as e:
        print(f"❌ 网络错误: {e}")
        return None


# --- 主程序 ---
def main():
    print("🎨 AI 绘图引擎启动...")

    # 模拟一个 Word 文档
    doc = Document()
    doc.add_heading('项目技术方案与进度', 0)

    req = "我们要建设一个电商网站，包含用户端、商家端、后台管理。工期3个月，第一个月做设计，第二个月开发，第三个月测试上线。"

    # --- 任务 A: 生成架构图 ---
    print("\n1️⃣ 正在生成【系统架构图】...")
    flow_code = generate_diagram_code(req, "flowchart")
    print(f"   代码: \n{flow_code[:50]}...")

    if flow_code:
        img_stream = mermaid_to_image(flow_code)
        if img_stream:
            doc.add_heading('1. 系统架构设计', level=1)
            doc.add_paragraph("下图展示了本项目的整体技术架构：")
            # 插入图片到 Word
            doc.add_picture(img_stream, width=Inches(6))
            print("   ✅ 架构图已插入 Word")

    # --- 任务 B: 生成甘特图 ---
    print("\n2️⃣ 正在生成【项目进度甘特图】...")
    gantt_code = generate_diagram_code(req, "gantt")
    print(f"   代码: \n{gantt_code[:50]}...")

    if gantt_code:
        img_stream = mermaid_to_image(gantt_code)
        if img_stream:
            doc.add_heading('2. 项目实施进度表', level=1)
            doc.add_paragraph("本项目详细实施计划如下：")
            doc.add_picture(img_stream, width=Inches(6))
            print("   ✅ 甘特图已插入 Word")

    # 保存
    doc.save("带图标书.docx")
    print("\n🎉 文件已保存：带图标书.docx")


if __name__ == "__main__":
    main()