import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO


# --- AI 核心函数 (增强版) ---

def ask_ai(messages, api_key):
    """
    通用 AI 请求函数
    messages: 完整的对话历史列表
    """
    url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    headers = {"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"}
    data = {
        "model": "glm-4",  # 建议用 glm-4，逻辑能力强
        "messages": messages,
        "temperature": 0.7
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            return f"Error: {response.text}"
    except Exception as e:
        return str(e)


def write_markdown_to_docx(doc, text):
    """
    增强版解析器：支持表格
    (简单实现：遇到 Markdown 表格就跳过或转纯文本，
    因为 python-docx 渲染表格比较复杂，这里先做文本转换)
    """
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue

        # 简单处理表格行 (以 | 开头)
        if line.startswith('|'):
            # 把表格行转成普通文本，避免 Word 格式乱掉
            # 进阶课可以学怎么用 python-docx 画真表格
            doc.add_paragraph(line, style='No Spacing')
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)


# --- 界面逻辑 ---

st.set_page_config(page_title="AI 标书专家 Pro", layout="wide")
st.title("🏆 AI 标书专家 Pro")

with st.sidebar:
    api_key = st.text_input("智谱 API Key", type="password")

col1, col2 = st.columns(2)
with col1:
    bid_req = st.text_area("1. 招标核心需求", height=200,
                           value="项目：智慧校园网络建设\n预算：500万\n工期：60天\n核心：全光网架构，覆盖教学楼、宿舍。")
with col2:
    company_info = st.text_area("2. 公司核心优势", height=200,
                                value="我司是华为金牌代理。拥有CCIE工程师5名。曾承建某大学万兆校园网项目。")

# 目录配置
default_chapters = """1. 项目背景与需求分析
2. 总体技术方案设计
3. 核心设备选型与参数
4. 项目实施进度计划
5. 售后服务与培训"""
chapters_input = st.text_area("3. 目录规划 (建议保持精简，让AI发挥)", value=default_chapters, height=150)

if st.button("🚀 启动深度生成", type="primary"):
    if not api_key:
        st.error("缺 API Key")
        st.stop()

    doc = Document()
    doc.add_heading('投标文件', 0)

    # --- 关键升级：全局记忆 ---
    # 这个变量用来存储关键决策，防止前后矛盾
    global_memory = f"""
    【全局约束】
    1. 必须严格遵守招标需求：{bid_req}
    2. 必须充分利用公司优势：{company_info}
    3. 全文语气必须统一，专业、严谨。
    """

    chapters = [c.strip() for c in chapters_input.split('\n') if c.strip()]

    progress = st.progress(0)
    status = st.empty()
    preview = st.container()

    for i, chapter in enumerate(chapters):
        status.info(f"正在深度撰写：{chapter} ...")

        doc.add_heading(chapter, level=1)

        # --- 步骤 1：思维链 (CoT) - 先列本章细纲 ---
        outline_prompt = f"""
        {global_memory}

        当前任务：请为【{chapter}】章节设计一个详细的二级子目录。
        要求：
        1. 只要子目录标题，不要内容。
        2. 逻辑要顺畅。
        3. 至少 3-4 个子标题。
        """
        # 这里用临时对话，不污染全局
        sub_outline = ask_ai([{"role": "user", "content": outline_prompt}], api_key)

        # --- 步骤 2：逐个子标题扩写 ---
        # 这一步是为了把字数撑起来，每章能写 1000+ 字

        full_chapter_content = ""

        # 解析子标题 (简单按行分割)
        sub_titles = [line for line in sub_outline.split('\n') if line.strip()]

        for sub_title in sub_titles:
            # 去掉可能的序号前缀
            clean_sub_title = sub_title.replace('#', '').strip()

            write_prompt = f"""
            {global_memory}

            当前任务：请撰写【{chapter}】下的子小节：【{clean_sub_title}】。

            要求：
            1. 内容详实，字数 300 字以上。
            2. 如果涉及到数据，请设计一个 Markdown 表格展示（如设备参数表、进度表）。
            3. 如果涉及到流程，请用 1. 2. 3. 步骤描述。
            4. 必须引用公司优势。
            """

            sub_content = ask_ai([{"role": "user", "content": write_prompt}], api_key)

            # 拼接到文档里
            doc.add_heading(clean_sub_title, level=2)
            write_markdown_to_docx(doc, sub_content)

            # 拼接到预览字符串
            full_chapter_content += f"### {clean_sub_title}\n{sub_content}\n\n"

        # --- 步骤 3：更新全局记忆 (可选) ---
        # 把这一章的核心摘要加回 global_memory，让后面章节知道前面写了啥
        # (为了省 Token，这里先省略这一步，进阶版可以做)

        # 预览
        with preview:
            with st.expander(f"✅ {chapter} (点击查看详情)", expanded=False):
                st.markdown(full_chapter_content)

        progress.progress((i + 1) / len(chapters))

    status.success("🎉 深度生成完成！")

    bio = BytesIO()
    doc.save(bio)
    st.download_button("📥 下载完整标书", bio.getvalue(), "深度标书.docx")